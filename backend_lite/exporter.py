"""
Court-Ready Exporter
====================

Generate DOCX and PDF exports for cross-examination plans with anchors.
"""

from typing import Any, Dict, List, Optional
from io import BytesIO


def _format_anchor(anchor: Dict[str, Any], doc_lookup: Dict[str, Any]) -> str:
    doc_id = anchor.get("doc_id")
    doc_name = None
    if doc_id and doc_id in doc_lookup:
        doc_name = getattr(doc_lookup[doc_id], "doc_name", None) or doc_id
    else:
        doc_name = doc_id or "מסמך לא ידוע"

    parts = [doc_name]
    if anchor.get("page_no") is not None:
        parts.append(f"עמ' {anchor.get('page_no')}")
    if anchor.get("paragraph_index") is not None:
        parts.append(f"פסקה {anchor.get('paragraph_index')}")
    if anchor.get("block_index") is not None and anchor.get("paragraph_index") is None:
        parts.append(f"בלוק {anchor.get('block_index')}")

    snippet = anchor.get("snippet") or ""
    if snippet:
        parts.append(f"\"{snippet}\"")

    return " | ".join(parts)


def build_cross_exam_docx(
    plan: Dict[str, Any],
    case_name: str,
    run_id: str,
    doc_lookup: Dict[str, Any],
) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX export") from exc

    doc = Document()

    def _add_right_paragraph(text: str):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return p

    title = doc.add_heading("תכנית חקירה נגדית", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    meta = doc.add_paragraph(f"תיק: {case_name} | הרצה: {run_id}")
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    case_settings = plan.get("case_settings") or {}
    heading = doc.add_heading("פרטי תיק", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    settings_lines = [
        ("מספר תיק", case_settings.get("case_number")),
        ("בית משפט", case_settings.get("court")),
        ("צד מיוצג", case_settings.get("our_side")),
        ("לקוח", case_settings.get("client_name")),
        ("צד שכנגד", case_settings.get("opponent_name")),
        ("סוג תיק", case_settings.get("case_type")),
        ("ערכאה", case_settings.get("court_level")),
        ("שפה", case_settings.get("language")),
    ]
    rendered = False
    for label, value in settings_lines:
        if value:
            _add_right_paragraph(f"{label}: {value}")
            rendered = True
    if not rendered:
        _add_right_paragraph("לא הוגדרו פרטי תיק נוספים.")

    ranked = plan.get("ranked_contradictions") or []
    heading = doc.add_heading("סתירות מדורגות", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if not ranked:
        _add_right_paragraph("לא נמצאו סתירות מדורגות לייצוא.")
    else:
        for idx, item in enumerate(ranked, start=1):
            scores = item.get("scores") or {}
            composite = scores.get("composite")
            composite_text = f"{composite:.2f}" if isinstance(composite, (int, float)) else "N/A"
            summary = f"{idx}. ציון משולב: {composite_text} | סוג: {item.get('type', '')}"
            if item.get("severity"):
                summary += f" | חומרה: {item.get('severity')}"
            if item.get("stage"):
                summary += f" | שלב: {item.get('stage')}"
            _add_right_paragraph(summary)
            if item.get("quote1"):
                _add_right_paragraph(f"ציטוט א': {item.get('quote1')}")
            if item.get("quote2"):
                _add_right_paragraph(f"ציטוט ב': {item.get('quote2')}")
            anchors = item.get("anchors") or []
            if anchors:
                _add_right_paragraph("עוגנים:")
                for anchor in anchors:
                    _add_right_paragraph(f"- {_format_anchor(anchor, doc_lookup)}")

    shifts = plan.get("version_shifts") or []
    heading = doc.add_heading("סטיות גרסה בעדויות", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if not shifts:
        _add_right_paragraph("לא נמצאו סטיות גרסה לייצוא.")
    else:
        for witness in shifts:
            _add_right_paragraph(f"עד: {witness.get('witness_name', 'לא ידוע')}")
            for shift in witness.get("shifts", []):
                _add_right_paragraph(f"- {shift.get('description', '')}")
                anchor_a = shift.get("anchor_a")
                anchor_b = shift.get("anchor_b")
                if anchor_a:
                    _add_right_paragraph(f"  עוגן א': {_format_anchor(anchor_a, doc_lookup)}")
                if anchor_b:
                    _add_right_paragraph(f"  עוגן ב': {_format_anchor(anchor_b, doc_lookup)}")

    heading = doc.add_heading("תכנית חקירה", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for stage in plan.get("stages", []):
        stage_heading = doc.add_heading(f"שלב {stage.get('stage', 'mid')}", level=2)
        stage_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for step in stage.get("steps", []):
            step_title = doc.add_paragraph(f"{step.get('title', '')} ({step.get('step_type', '')})")
            step_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            question = doc.add_paragraph(step.get("question", ""))
            question.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if step.get("do_not_ask_flag"):
                warning = doc.add_paragraph("DON'T ASK THIS: " + (step.get("do_not_ask_reason") or "סיכון גבוה."))
                warning.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            anchors = step.get("anchors") or []
            if anchors:
                doc.add_paragraph("עוגנים:", style=None).alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for anchor in anchors:
                    line = _format_anchor(anchor, doc_lookup)
                    para = doc.add_paragraph(f"- {line}")
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            branches = step.get("branches") or []
            if branches:
                doc.add_paragraph("הסתעפויות:", style=None).alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for branch in branches:
                    trigger = branch.get("trigger", "")
                    para = doc.add_paragraph(f"* {trigger}")
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for follow_up in branch.get("follow_up_questions", []):
                        f_para = doc.add_paragraph(f"  - {follow_up}")
                        f_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    appendix = plan.get("appendix_anchors") or []
    heading = doc.add_heading("נספח: קטעי ראיות", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if not appendix:
        _add_right_paragraph("לא נמצאו קטעי ראיות להצגה.")
    else:
        for anchor in appendix:
            _add_right_paragraph(_format_anchor(anchor, doc_lookup))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_cross_exam_pdf(
    plan: Dict[str, Any],
    case_name: str,
    run_id: str,
    doc_lookup: Dict[str, Any],
) -> bytes:
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError as exc:
        raise RuntimeError("reportlab is required for PDF export") from exc

    # Register font with Hebrew support
    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    pdfmetrics.registerFont(TTFont("DejaVuSans", font_path))

    try:
        from bidi.algorithm import get_display
    except Exception:
        def get_display(value: str) -> str:
            return value

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setFont("DejaVuSans", 14)

    width, height = A4
    y = height - 50

    def draw_text(text: str, size: int = 12):
        nonlocal y
        if y < 80:
            c.showPage()
            c.setFont("DejaVuSans", size)
            y = height - 50
        c.setFont("DejaVuSans", size)
        c.drawRightString(width - 40, y, get_display(text))
        y -= size + 6

    draw_text("תכנית חקירה נגדית", 16)
    draw_text(f"תיק: {case_name} | הרצה: {run_id}", 11)

    case_settings = plan.get("case_settings") or {}
    draw_text("פרטי תיק", 14)
    settings_lines = [
        ("מספר תיק", case_settings.get("case_number")),
        ("בית משפט", case_settings.get("court")),
        ("צד מיוצג", case_settings.get("our_side")),
        ("לקוח", case_settings.get("client_name")),
        ("צד שכנגד", case_settings.get("opponent_name")),
        ("סוג תיק", case_settings.get("case_type")),
        ("ערכאה", case_settings.get("court_level")),
        ("שפה", case_settings.get("language")),
    ]
    rendered = False
    for label, value in settings_lines:
        if value:
            draw_text(f"{label}: {value}", 11)
            rendered = True
    if not rendered:
        draw_text("לא הוגדרו פרטי תיק נוספים.", 11)

    ranked = plan.get("ranked_contradictions") or []
    draw_text("סתירות מדורגות", 14)
    if not ranked:
        draw_text("לא נמצאו סתירות מדורגות לייצוא.", 11)
    else:
        for idx, item in enumerate(ranked, start=1):
            scores = item.get("scores") or {}
            composite = scores.get("composite")
            composite_text = f"{composite:.2f}" if isinstance(composite, (int, float)) else "N/A"
            summary = f"{idx}. ציון משולב: {composite_text} | סוג: {item.get('type', '')}"
            if item.get("severity"):
                summary += f" | חומרה: {item.get('severity')}"
            if item.get("stage"):
                summary += f" | שלב: {item.get('stage')}"
            draw_text(summary, 11)
            if item.get("quote1"):
                draw_text(f"ציטוט א': {item.get('quote1')}", 10)
            if item.get("quote2"):
                draw_text(f"ציטוט ב': {item.get('quote2')}", 10)
            anchors = item.get("anchors") or []
            for anchor in anchors:
                draw_text(_format_anchor(anchor, doc_lookup), 10)

    shifts = plan.get("version_shifts") or []
    draw_text("סטיות גרסה בעדויות", 14)
    if not shifts:
        draw_text("לא נמצאו סטיות גרסה לייצוא.", 11)
    else:
        for witness in shifts:
            draw_text(f"עד: {witness.get('witness_name', 'לא ידוע')}", 11)
            for shift in witness.get("shifts", []):
                draw_text(f"- {shift.get('description', '')}", 10)
                anchor_a = shift.get("anchor_a")
                anchor_b = shift.get("anchor_b")
                if anchor_a:
                    draw_text(f"  עוגן א': {_format_anchor(anchor_a, doc_lookup)}", 9)
                if anchor_b:
                    draw_text(f"  עוגן ב': {_format_anchor(anchor_b, doc_lookup)}", 9)

    draw_text("תכנית חקירה", 14)
    for stage in plan.get("stages", []):
        draw_text(f"שלב {stage.get('stage', 'mid')}", 13)
        for step in stage.get("steps", []):
            draw_text(f"{step.get('title', '')} ({step.get('step_type', '')})", 12)
            draw_text(step.get("question", ""), 11)
            if step.get("do_not_ask_flag"):
                draw_text("DON'T ASK THIS: " + (step.get("do_not_ask_reason") or "סיכון גבוה."), 10)

            anchors = step.get("anchors") or []
            for anchor in anchors:
                draw_text(_format_anchor(anchor, doc_lookup), 10)

            branches = step.get("branches") or []
            for branch in branches:
                draw_text(f"הסתעפות: {branch.get('trigger', '')}", 10)
                for follow_up in branch.get("follow_up_questions", []):
                    draw_text(f"- {follow_up}", 10)

    appendix = plan.get("appendix_anchors") or []
    draw_text("נספח: קטעי ראיות", 14)
    if not appendix:
        draw_text("לא נמצאו קטעי ראיות להצגה.", 11)
    else:
        for anchor in appendix:
            draw_text(_format_anchor(anchor, doc_lookup), 10)

    c.save()
    buf.seek(0)
    return buf.read()
