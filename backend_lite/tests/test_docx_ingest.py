"""
DOCX Ingestion Tests (Epic A1)
"""

import json
import zipfile
from pathlib import Path

import pytest

from backend_lite.ingest.docx import DOCXParser
from backend_lite.ingest.base import ParserError

FIXTURES_DIR = Path(__file__).parent / "fixtures" / "docx"


def _build_docx_with_paragraphs(paragraphs, tmp_path: Path) -> bytes:
    from docx import Document

    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    file_path = tmp_path / "simple.docx"
    doc.save(file_path)
    return file_path.read_bytes()


def _build_docx_with_table(rows, tmp_path: Path) -> bytes:
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    file_path = tmp_path / "table.docx"
    doc.save(file_path)
    return file_path.read_bytes()


def _build_docx_with_document_xml(xml_text: str, tmp_path: Path) -> bytes:
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""
    file_path = tmp_path / "track.docx"
    with zipfile.ZipFile(file_path, "w") as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("word/document.xml", xml_text)
    return file_path.read_bytes()


def test_ingest_docx_ok(tmp_path: Path):
    paragraphs = FIXTURES_DIR.joinpath("simple_hebrew.txt").read_text(encoding="utf-8").splitlines()
    data = _build_docx_with_paragraphs(paragraphs, tmp_path)

    result = DOCXParser().parse(data, filename="simple.docx")
    assert result.full_text
    assert result.page_count == 1
    assert len(result.pages[0].blocks) >= 2

    for block in result.pages[0].blocks:
        if block.char_start is None or block.char_end is None:
            continue
        assert result.full_text[block.char_start:block.char_end] == block.text


def test_ingest_docx_table_ok(tmp_path: Path):
    rows = json.loads(FIXTURES_DIR.joinpath("table_hebrew.json").read_text(encoding="utf-8"))["rows"]
    data = _build_docx_with_table(rows, tmp_path)

    result = DOCXParser().parse(data, filename="table.docx")
    assert result.full_text
    assert result.metadata.get("table_count") == 1
    assert any(" | ".join(rows[0]) in block.text for block in result.pages[0].blocks)


def test_ingest_docx_track_changes_returns_user_error(tmp_path: Path):
    xml_text = FIXTURES_DIR.joinpath("track_changes_document.xml").read_text(encoding="utf-8")
    data = _build_docx_with_document_xml(xml_text, tmp_path)

    with pytest.raises(ParserError) as exc:
        DOCXParser().parse(data, filename="track_changes.docx")

    assert exc.value.code == "docx_track_changes"
