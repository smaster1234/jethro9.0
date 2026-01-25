"""
DOCX Parser
===========

Microsoft Word document parser using python-docx.
"""

import io
import os
import logging
import zipfile
from typing import List, Optional, Tuple
from xml.etree import ElementTree as ET

from .base import (
    DocumentParser,
    ParseResult,
    PageContent,
    BlockContent,
    ParserError,
    DocxTrackChangesError,
    normalize_text
)

DOCX_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
}

logger = logging.getLogger(__name__)
_DOCX_DEBUG = os.environ.get("DOCX_INGEST_DEBUG", "").strip().lower() in ("1", "true", "yes")


def _debug_log(code: str, exc: Exception) -> None:
    if _DOCX_DEBUG:
        logger.debug("docx_ingest_error code=%s exc=%s", code, exc.__class__.__name__)


def _read_document_xml(data: bytes) -> Optional[str]:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            if "word/document.xml" not in zf.namelist():
                return None
            xml_bytes = zf.read("word/document.xml")
            return xml_bytes.decode("utf-8", errors="ignore")
    except zipfile.BadZipFile:
        return None


def _contains_track_changes(root: ET.Element) -> bool:
    for tag in ("ins", "del", "moveFrom", "moveTo"):
        if root.findall(f".//w:{tag}", DOCX_NS):
            return True
    return False


def _extract_text_from_paragraph(node: ET.Element) -> str:
    parts: List[str] = []
    for text_node in node.findall(".//w:t", DOCX_NS):
        if text_node.text:
            parts.append(text_node.text)
    return "".join(parts).strip()


def _extract_blocks_from_xml(xml_text: str) -> Tuple[List[Tuple[str, Optional[int]]], int, int]:
    root = ET.fromstring(xml_text)
    if _contains_track_changes(root):
        raise DocxTrackChangesError()

    body = root.find("w:body", DOCX_NS)
    if body is None:
        return [], 0, 0

    blocks: List[Tuple[str, Optional[int]]] = []
    paragraph_index = 0
    table_count = 0

    for child in list(body):
        if child.tag == f"{{{DOCX_NS['w']}}}p":
            text = _extract_text_from_paragraph(child)
            if text:
                blocks.append((text, paragraph_index))
            paragraph_index += 1
        elif child.tag == f"{{{DOCX_NS['w']}}}tbl":
            table_count += 1
            for row in child.findall(".//w:tr", DOCX_NS):
                row_cells: List[str] = []
                for cell in row.findall(".//w:tc", DOCX_NS):
                    cell_parts: List[str] = []
                    for para in cell.findall(".//w:p", DOCX_NS):
                        cell_text = _extract_text_from_paragraph(para)
                        if cell_text:
                            cell_parts.append(cell_text)
                    cell_text = " ".join(cell_parts).strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    blocks.append((" | ".join(row_cells), None))

    return blocks, paragraph_index, table_count


def _build_result(block_items: List[Tuple[str, Optional[int]]]) -> Tuple[List[BlockContent], str]:
    blocks: List[BlockContent] = []
    full_text_parts: List[str] = []
    char_offset = 0

    for text, paragraph_index in block_items:
        normalized = normalize_text(text)
        if not normalized:
            continue

        blocks.append(BlockContent(
            text=normalized,
            block_index=len(blocks),
            page_no=1,
            paragraph_index=paragraph_index,
            char_start=char_offset,
            char_end=char_offset + len(normalized)
        ))

        full_text_parts.append(normalized)
        char_offset += len(normalized) + 1

    full_text = "\n".join(full_text_parts)
    return blocks, full_text


class DOCXParser(DocumentParser):
    """
    Microsoft Word (.docx) parser.

    Uses python-docx to extract text with paragraph structure.
    """

    @property
    def supported_mimes(self) -> List[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"  # .doc files (limited support)
        ]

    def parse(self, data: bytes, filename: str = None) -> ParseResult:
        """Parse DOCX file"""
        try:
            from docx import Document
        except ImportError as exc:
            _debug_log("docx_missing_dependency", exc)
            raise ParserError(
                "python-docx is required for DOCX parsing.",
                code="docx_missing_dependency",
                user_message="נדרש רכיב python-docx לעיבוד DOCX. התקן תלות זו."
            )

        document_xml = _read_document_xml(data)
        if document_xml:
            try:
                _extract_blocks_from_xml(document_xml)
            except DocxTrackChangesError as exc:
                _debug_log("docx_track_changes", exc)
                raise
            except Exception as exc:
                _debug_log("docx_xml_parse_failed", exc)
                pass

        try:
            # Load document from bytes
            doc = Document(io.BytesIO(data))

            block_items: List[Tuple[str, Optional[int]]] = []

            # Process paragraphs
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                block_items.append((text, idx))

            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)

                    if row_texts:
                        block_items.append((" | ".join(row_texts), None))

            blocks, full_text = _build_result(block_items)

            # Create single page (DOCX doesn't have native page breaks)
            page = PageContent(
                page_no=1,
                text=full_text,
                blocks=blocks,
                char_start=0,
                char_end=len(full_text)
            )

            # Extract metadata
            metadata = {}
            try:
                core_props = doc.core_properties
                if core_props.author:
                    metadata['author'] = core_props.author
                if core_props.title:
                    metadata['title'] = core_props.title
                if core_props.created:
                    metadata['created'] = core_props.created.isoformat()
                if core_props.modified:
                    metadata['modified'] = core_props.modified.isoformat()
            except:
                pass

            metadata['paragraph_count'] = len(blocks)
            metadata['table_count'] = len(doc.tables)

            return ParseResult(
                full_text=full_text,
                pages=[page],
                page_count=1,
                language=self._detect_language(full_text),
                metadata=metadata
            )

        except DocxTrackChangesError as exc:
            _debug_log("docx_track_changes", exc)
            raise
        except Exception as exc:
            if document_xml:
                try:
                    block_items, paragraph_count, table_count = _extract_blocks_from_xml(document_xml)
                    blocks, full_text = _build_result(block_items)
                    page = PageContent(
                        page_no=1,
                        text=full_text,
                        blocks=blocks,
                        char_start=0,
                        char_end=len(full_text)
                    )
                    return ParseResult(
                        full_text=full_text,
                        pages=[page],
                        page_count=1,
                        language=self._detect_language(full_text),
                        metadata={
                            "paragraph_count": paragraph_count,
                            "table_count": table_count,
                            "parser": "xml_fallback",
                        }
                    )
                except DocxTrackChangesError as exc:
                    _debug_log("docx_track_changes", exc)
                    raise
                except Exception as inner_exc:
                    _debug_log("docx_xml_parse_failed", inner_exc)
                    pass

            _debug_log("docx_parse_failed", exc)
            raise ParserError(
                "Failed to parse DOCX file.",
                code="docx_parse_failed",
                user_message="המסמך לא ניתן לעיבוד. נסה לשמור מחדש כ‑DOCX תקין.",
            )

    def _detect_language(self, text: str) -> str:
        """Simple language detection"""
        if not text:
            return "unknown"

        hebrew_chars = sum(1 for c in text if '\u0590' <= c <= '\u05FF')
        total_alpha = sum(1 for c in text if c.isalpha())

        if total_alpha == 0:
            return "unknown"

        return "he" if hebrew_chars / total_alpha > 0.3 else "en"
